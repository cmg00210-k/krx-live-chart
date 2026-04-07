"""
CheeseStock ANATOMY V7 -- 4-Deliverable Document Generator
Hybrid pipeline: pandoc (MD->DOCX) + python-docx (post-processing)
PDF pipeline:    pandoc (MD->XeLaTeX->PDF) with custom template

Deliverable 1: Executive Brief (~20-25 pages)
Deliverable 2: Technical Architecture (~120-150 pages)
Deliverable 3: Theoretical Foundations & Appendices (~200+ pages)
Deliverable 4: Structure Flow (~16-18 pages PDF / ~25 pages DOCX)

Usage:
    python scripts/generate_word.py           # Generate all 4 (DOCX)
    python scripts/generate_word.py --exec    # Deliverable 1 only
    python scripts/generate_word.py --tech    # Deliverable 2 only
    python scripts/generate_word.py --theory  # Deliverable 3 only
    python scripts/generate_word.py --d4      # Deliverable 4 only
    python scripts/generate_word.py --pdf     # All targets as PDF
    python scripts/generate_word.py --d4 --pdf  # D4 as PDF
"""

import subprocess
import sys
import os
from pathlib import Path

# -- Configuration -----------------------------------------------------------
PROJECT_ROOT = Path(__file__).resolve().parent.parent
ANATOMY_DIR = PROJECT_ROOT / "docs" / "anatomy"
OUTPUT_DIR = PROJECT_ROOT / "docs"
REFERENCE_FILE = OUTPUT_DIR / "reference.docx"

# 3-Deliverable output files
DELIVERABLES = {
    "exec": {
        "output": OUTPUT_DIR / "CheeseStock_D1_Executive_Brief.docx",
        "title": "Deliverable 1: Executive Brief",
        "files": [
            "deliverable1_executive/P0_executive_summary.md",
        ],
    },
    "tech": {
        "output": OUTPUT_DIR / "CheeseStock_D2_Technical_Architecture.docx",
        "title": "Deliverable 2: Technical Architecture",
        "files": [
            # Part 1: Platform Architecture & Data Infrastructure
            "S1_api_pipeline_v7_sec1to4.md",
            "S1_api_pipeline_v7_sec5to8.md",
            "S1_api_pipeline_v7_sec9.md",
            "S5_lifecycle_workers_v7.md",
            # Part 2: Analytical Engine
            "S3_ta_methods_v7.md",
            "S3_signal_backtester_v7.md",
            "S3_confidence_chain_v7.md",
            # Part 3: Validation & Risk Management (NEW)
            "deliverable2_technical/P3_validation_risk.md",
            # Part 4: Rendering & User Experience
            "S4_chart_rendering_v7.md",
            "S5_ui_architecture_v7.md",
            # Part 5: Theory Summary Cards (NEW)
            "deliverable2_technical/P5_theory_summary_cards.md",
        ],
    },
    "theory": {
        "output": OUTPUT_DIR / "CheeseStock_D3_Theoretical_Foundations.docx",
        "title": "Deliverable 3: Theoretical Foundations & Appendices",
        "files": [
            # Full theory texts
            "S2_theoretical_basis_v7.md",
            "S2_sec23_finance_behavioral_v7.md",
            "S2_sec25_macroeconomics_v7.md",
            "S2_sec26_microeconomics_v7.md",
            "S2_sec27_derivatives_v7.md",
            "S2_sec28_bonds_credit_v7.md",
            "S2_sec29_rl_game_control_v7.md",
            # Appendices from V5 archive
            "archive_v5/S2_formula_appendix.md",
            "archive_v5/S2_constant_registry.md",
            "archive_v5/S3_backtest_methodology.md",
            "archive_v5/S4_canvas_layer_spec.md",
            # Verification (moved to appendix)
            "S0_index_v7.md",
            "S0_consistency_audit_v7.md",
            "S0_cross_stage_verification_v7.md",
            "S0_formula_fidelity_v7.md",
        ],
    },
    "d4": {
        "output": OUTPUT_DIR / "CheeseStock_D4_Structure_Flow.docx",
        "title": "Deliverable 4: Structure Flow",
        "files": [
            "deliverable4_structure_flow/D4_structure_flow.md",
        ],
    },
    "d4ko": {
        "output": OUTPUT_DIR / "CheeseStock_D4_Structure_Flow_KO.docx",
        "title": "CheeseStock ANATOMY V7 -- Structure Flow (KO)",
        "files": [
            "deliverable4_structure_flow/D4_structure_flow_ko.md",
        ],
    },
}

# Colors
NAVY = "2C3E6B"
GOLD = "C9A84C"
DARK = "1A1A2E"
BODY_TEXT_COLOR = "1A1A1A"
MUTED = "555555"
TABLE_ODD_BG = "F7F8FC"
CODE_BG = "F4F4F6"

# Fonts
FONT_KR = "Malgun Gothic"
FONT_CODE = "JetBrains Mono"
FONT_LATIN = "Segoe UI"

# Minimum lines for H2 page break (sparse sections stay on same page)
MIN_H2_LINES_FOR_BREAK = 15

# PDF pipeline paths
TEMPLATE_DIR = PROJECT_ROOT / "scripts" / "templates"
PDF_TEMPLATE = TEMPLATE_DIR / "cheesestock-d4.tex"
LUA_FILTER_DIAGRAM = TEMPLATE_DIR / "diagram-protect.lua"

# PDF output mapping
PDF_OUTPUTS = {
    "exec": OUTPUT_DIR / "CheeseStock_D1_Executive_Brief.pdf",
    "tech": OUTPUT_DIR / "CheeseStock_D2_Technical_Architecture.pdf",
    "theory": OUTPUT_DIR / "CheeseStock_D3_Theoretical_Foundations.pdf",
    "d4": OUTPUT_DIR / "CheeseStock_D4_Structure_Flow.pdf",
    "d4ko": OUTPUT_DIR / "CheeseStock_ANATOMY_V7_KO.pdf",
}


# -- Step 1: Create reference.docx template ----------------------------------
def create_reference_template():
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.oxml.ns import qn, nsdecls
    from docx.oxml import parse_xml

    doc = Document()
    for section in doc.sections:
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
        section.left_margin = Cm(3.0)
        section.right_margin = Cm(2.5)
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)

    def setup_style(style, font_name, size_pt, bold=False, color_hex=None,
                    space_before=0, space_after=0, keep_with_next=False):
        font = style.font
        font.name = font_name
        font.size = Pt(size_pt)
        font.bold = bold
        if color_hex:
            font.color.rgb = RGBColor.from_string(color_hex)
        rpr = style.element.get_or_add_rPr()
        ea = rpr.find(qn('w:rFonts'))
        if ea is None:
            ea = parse_xml(f'<w:rFonts {nsdecls("w")} w:eastAsia="{FONT_KR}"/>')
            rpr.insert(0, ea)
        else:
            ea.set(qn('w:eastAsia'), FONT_KR)
        pf = style.paragraph_format
        pf.space_before = Pt(space_before)
        pf.space_after = Pt(space_after)
        pf.keep_with_next = keep_with_next
        pf.widow_control = True

    styles = doc.styles
    setup_style(styles['Normal'], FONT_LATIN, 10, color_hex=BODY_TEXT_COLOR, space_after=6)

    h1 = styles['Heading 1']
    setup_style(h1, FONT_LATIN, 24, bold=True, color_hex=DARK,
                space_before=36, space_after=12, keep_with_next=True)
    h1.paragraph_format.page_break_before = True

    h2 = styles['Heading 2']
    setup_style(h2, FONT_LATIN, 18, bold=True, color_hex=DARK,
                space_before=30, space_after=10, keep_with_next=True)

    h3 = styles['Heading 3']
    setup_style(h3, FONT_LATIN, 13, bold=True, color_hex=NAVY,
                space_before=18, space_after=6, keep_with_next=True)

    try:
        h4 = styles['Heading 4']
    except KeyError:
        h4 = styles.add_style('Heading 4', 1)
    setup_style(h4, FONT_LATIN, 10.5, bold=True, color_hex=MUTED,
                space_before=12, space_after=4, keep_with_next=True)

    for sn in ['Heading 1', 'Heading 2', 'Heading 3', 'Heading 4']:
        doc.add_paragraph('', style=sn)
    doc.add_paragraph('', style='Normal')

    doc.save(str(REFERENCE_FILE))
    print(f"  [OK] Reference template: {REFERENCE_FILE}")


# -- Step 2: Run pandoc -------------------------------------------------------
def find_pandoc():
    pandoc_paths = [
        r"C:\Users\seth1\AppData\Local\Microsoft\WinGet\Packages\JohnMacFarlane.Pandoc_Microsoft.Winget.Source_8wekyb3d8bbwe\pandoc-3.9.0.2\pandoc.exe",
        r"C:\Users\seth1\AppData\Local\Pandoc\pandoc.exe",
        r"C:\Program Files\Pandoc\pandoc.exe",
    ]
    for p in pandoc_paths:
        if os.path.exists(p):
            return p
    return "pandoc"


def build_file_list(file_names):
    ordered = []
    for f in file_names:
        p = ANATOMY_DIR / f
        if p.exists():
            ordered.append(p)
        else:
            print(f"  [WARN] Missing: {f}")
    return ordered


def run_pandoc(file_list, output_file):
    pandoc = find_pandoc()
    cmd = [
        pandoc, "--from=markdown", "--to=docx",
        f"--reference-doc={REFERENCE_FILE}",
        "--toc", "--toc-depth=3", "--standalone",
        f"--output={output_file}", "--wrap=none",
    ]
    cmd.extend([str(f) for f in file_list])
    result = subprocess.run(cmd, capture_output=True, text=True, encoding='utf-8')
    if result.returncode != 0:
        print(f"  [FAIL] pandoc error:\n{result.stderr}")
        return False
    return True


# -- Step 3: Post-process with python-docx ------------------------------------
def post_process(output_file):
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.oxml.ns import qn, nsdecls
    from docx.oxml import parse_xml
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document(str(output_file))

    # Page setup + footer
    for section in doc.sections:
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
        section.left_margin = Cm(3.0)
        section.right_margin = Cm(2.5)
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        footer = section.footer
        footer.is_linked_to_previous = False
        fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _add_page_number(fp)

    # Build H2 content-length map for smart page breaks
    h2_content_lengths = _measure_h2_lengths(doc.paragraphs)

    # Style paragraphs
    para_list = doc.paragraphs
    for i, para in enumerate(para_list):
        _style_paragraph(para, i, para_list, h2_content_lengths)

    # Style tables
    for table in doc.tables:
        _style_table(table)

    # Wide tables
    _handle_wide_tables(doc)

    # Visual separators
    _add_visual_separators(doc)

    # Korean fonts
    _fix_korean_fonts(doc)

    doc.save(str(output_file))
    return _get_stats(doc)


def _measure_h2_lengths(para_list):
    """Count paragraphs between each H2 and next H1/H2 to determine content density."""
    h2_indices = []
    for i, p in enumerate(para_list):
        if p.style and p.style.name == "Heading 2":
            h2_indices.append(i)

    lengths = {}
    for idx in h2_indices:
        count = 0
        for j in range(idx + 1, len(para_list)):
            sn = para_list[j].style.name if para_list[j].style else ""
            if sn in ("Heading 1", "Heading 2"):
                break
            if para_list[j].text.strip():
                count += 1
        lengths[idx] = count
    return lengths


def _add_page_number(paragraph):
    from docx.oxml.ns import nsdecls
    from docx.oxml import parse_xml
    from docx.shared import Pt, RGBColor

    run = paragraph.add_run("\u2014 ")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor.from_string(MUTED)

    r2 = paragraph.add_run()
    r2._r.append(parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>'))
    r3 = paragraph.add_run()
    r3._r.append(parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>'))
    r4 = paragraph.add_run()
    r4._r.append(parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>'))

    r5 = paragraph.add_run(" \u2014")
    r5.font.size = Pt(9)
    r5.font.color.rgb = RGBColor.from_string(MUTED)


def _style_paragraph(para, idx, para_list, h2_lengths):
    from docx.shared import Pt, RGBColor
    from docx.oxml.ns import qn, nsdecls
    from docx.oxml import parse_xml

    style_name = para.style.name if para.style else ""

    if style_name.startswith("Heading"):
        para.paragraph_format.keep_with_next = True
        para.paragraph_format.widow_control = True

        if style_name == "Heading 1":
            para.paragraph_format.page_break_before = True
            para.paragraph_format.space_before = Pt(36)
            para.paragraph_format.space_after = Pt(12)
            for run in para.runs:
                run.font.size = Pt(24)
                run.font.bold = True
                run.font.color.rgb = RGBColor.from_string(DARK)

        elif style_name == "Heading 2":
            # Smart page break: skip if after H1 or if section is sparse
            should_break = True

            # Skip if immediately after H1
            for j in range(idx - 1, max(idx - 4, -1), -1):
                prev = para_list[j]
                if prev.style and prev.style.name == "Heading 1":
                    should_break = False
                    break
                elif prev.text.strip():
                    break

            # Skip if section content < MIN_H2_LINES_FOR_BREAK
            content_len = h2_lengths.get(idx, 999)
            if content_len < MIN_H2_LINES_FOR_BREAK:
                should_break = False

            para.paragraph_format.page_break_before = should_break
            para.paragraph_format.space_before = Pt(30)
            para.paragraph_format.space_after = Pt(10)
            for run in para.runs:
                run.font.size = Pt(18)
                run.font.bold = True
                run.font.color.rgb = RGBColor.from_string(DARK)

        elif style_name == "Heading 3":
            para.paragraph_format.space_before = Pt(18)
            para.paragraph_format.space_after = Pt(6)
            for run in para.runs:
                run.font.size = Pt(13)
                run.font.bold = True
                run.font.color.rgb = RGBColor.from_string(NAVY)

        elif style_name == "Heading 4":
            para.paragraph_format.space_before = Pt(12)
            para.paragraph_format.space_after = Pt(4)
            for run in para.runs:
                run.font.size = Pt(10.5)
                run.font.bold = True
                run.font.color.rgb = RGBColor.from_string(MUTED)

    elif any(kw in style_name for kw in
             ("Source Code", "Verbatim", "Code Block", "Preformatted")):
        # Detect wide ASCII diagrams (>60 chars) → use smaller font for alignment
        max_line_len = max((len(r.text) for r in para.runs if r.text), default=0)
        is_wide = max_line_len > 60
        code_size = Pt(7.5) if is_wide else Pt(8.5)
        code_indent = Pt(6) if is_wide else Pt(12)

        para.paragraph_format.space_before = Pt(8)
        para.paragraph_format.space_after = Pt(8)
        para.paragraph_format.keep_together = True
        para.paragraph_format.left_indent = code_indent
        para.paragraph_format.right_indent = code_indent
        for run in para.runs:
            run.font.name = FONT_CODE
            run.font.size = code_size
            # Ensure eastAsia also uses monospace for box-drawing chars
            rpr = run._r.get_or_add_rPr()
            rf = rpr.find(qn('w:rFonts'))
            if rf is None:
                rpr.insert(0, parse_xml(
                    f'<w:rFonts {nsdecls("w")} w:ascii="{FONT_CODE}" '
                    f'w:hAnsi="{FONT_CODE}" w:eastAsia="{FONT_CODE}" '
                    f'w:cs="{FONT_CODE}"/>'))
            else:
                rf.set(qn('w:ascii'), FONT_CODE)
                rf.set(qn('w:hAnsi'), FONT_CODE)
                rf.set(qn('w:eastAsia'), FONT_CODE)
                rf.set(qn('w:cs'), FONT_CODE)
        pPr = para._p.get_or_add_pPr()
        existing = pPr.find(qn('w:shd'))
        if existing is not None:
            pPr.remove(existing)
        pPr.append(parse_xml(f'<w:shd {nsdecls("w")} w:fill="{CODE_BG}" w:val="clear"/>'))

    elif style_name == "Body Text":
        para.paragraph_format.space_before = Pt(6)
        para.paragraph_format.space_after = Pt(6)
        para.paragraph_format.left_indent = Pt(18)
        for run in para.runs:
            run.font.color.rgb = RGBColor.from_string(MUTED)
            run.font.size = Pt(9.5)
            run.font.italic = True
        pPr = para._p.get_or_add_pPr()
        existing = pPr.find(qn('w:pBdr'))
        if existing is not None:
            pPr.remove(existing)
        pPr.append(parse_xml(
            f'<w:pBdr {nsdecls("w")}>'
            f'  <w:left w:val="single" w:sz="18" w:space="6" w:color="{GOLD}"/>'
            f'</w:pBdr>'
        ))

    else:
        para.paragraph_format.widow_control = True
        para.paragraph_format.space_after = Pt(6)


def _style_table(table):
    from docx.shared import Pt, RGBColor
    from docx.oxml.ns import qn, nsdecls
    from docx.oxml import parse_xml
    from docx.enum.table import WD_TABLE_ALIGNMENT

    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    tblW = tblPr.find(qn('w:tblW'))
    if tblW is None:
        tblPr.append(parse_xml(f'<w:tblW {nsdecls("w")} w:w="5000" w:type="pct"/>'))
    else:
        tblW.set(qn('w:w'), '5000')
        tblW.set(qn('w:type'), 'pct')

    for row_idx, row in enumerate(table.rows):
        for cell in row.cells:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            if row_idx == 0:
                shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{NAVY}" w:val="clear"/>')
                existing = tcPr.find(qn('w:shd'))
                if existing is not None:
                    tcPr.remove(existing)
                tcPr.append(shd)
                for p in cell.paragraphs:
                    for r in p.runs:
                        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                        r.font.size = Pt(9)
                        r.font.bold = True
            else:
                bg = TABLE_ODD_BG if row_idx % 2 == 1 else "FFFFFF"
                shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{bg}" w:val="clear"/>')
                existing = tcPr.find(qn('w:shd'))
                if existing is not None:
                    tcPr.remove(existing)
                tcPr.append(shd)
                for p in cell.paragraphs:
                    for r in p.runs:
                        r.font.size = Pt(9)
                        r.font.color.rgb = RGBColor.from_string(BODY_TEXT_COLOR)
            _set_cell_padding(tcPr, top=40, bottom=40, left=80, right=80)

    if len(table.rows) > 0:
        trPr = table.rows[0]._tr.get_or_add_trPr()
        if trPr.find(qn('w:tblHeader')) is None:
            trPr.append(parse_xml(f'<w:tblHeader {nsdecls("w")}/>'))


def _set_cell_padding(tcPr, top=0, bottom=0, left=0, right=0):
    from docx.oxml.ns import qn, nsdecls
    from docx.oxml import parse_xml
    mar = tcPr.find(qn('w:tcMar'))
    if mar is None:
        mar = parse_xml(f'<w:tcMar {nsdecls("w")}/>')
        tcPr.append(mar)
    for side, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        el = mar.find(qn(f'w:{side}'))
        if el is None:
            mar.append(parse_xml(f'<w:{side} {nsdecls("w")} w:w="{val}" w:type="dxa"/>'))
        else:
            el.set(qn('w:w'), str(val))
            el.set(qn('w:type'), 'dxa')


def _handle_wide_tables(doc):
    from docx.shared import Pt
    from docx.oxml.ns import qn, nsdecls
    from docx.oxml import parse_xml

    landscape_count = 0
    for table in doc.tables:
        ncols = len(table.columns)
        if ncols < 7:
            continue
        font_size = Pt(7) if ncols >= 9 else Pt(8)
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for r in p.runs:
                        r.font.size = font_size

    body = doc.element.body
    for tbl in reversed(list(body.findall(qn('w:tbl')))):
        first_row = tbl.find(qn('w:tr'))
        if first_row is None:
            continue
        if len(first_row.findall(qn('w:tc'))) >= 9:
            landscape_count += 1
            tbl.addnext(parse_xml(
                f'<w:p {nsdecls("w")}><w:pPr><w:sectPr>'
                f'<w:pgSz w:w="11906" w:h="16838" w:orient="portrait"/>'
                f'<w:pgMar w:top="1418" w:right="1418" w:bottom="1418" w:left="1701" w:header="720" w:footer="720"/>'
                f'<w:type w:val="continuous"/>'
                f'</w:sectPr></w:pPr></w:p>'
            ))
            tbl.addprevious(parse_xml(
                f'<w:p {nsdecls("w")}><w:pPr><w:sectPr>'
                f'<w:pgSz w:w="16838" w:h="11906" w:orient="landscape"/>'
                f'<w:pgMar w:top="1418" w:right="1701" w:bottom="1418" w:left="1418" w:header="720" w:footer="720"/>'
                f'<w:type w:val="continuous"/>'
                f'</w:sectPr></w:pPr></w:p>'
            ))
    if landscape_count:
        print(f"    {landscape_count} landscape table(s)")


def _add_visual_separators(doc):
    from docx.oxml.ns import nsdecls
    from docx.oxml import parse_xml

    h1_elems = []
    h2_elems = []
    for para in doc.paragraphs:
        if para.style and para.style.name == "Heading 1":
            h1_elems.append(para._p)
        elif para.style and para.style.name == "Heading 2":
            h2_elems.append(para._p)

    for p in h1_elems:
        p.addnext(parse_xml(
            f'<w:p {nsdecls("w")}><w:pPr>'
            f'<w:pBdr><w:bottom w:val="single" w:sz="12" w:space="1" w:color="{GOLD}"/></w:pBdr>'
            f'<w:spacing w:before="0" w:after="120"/>'
            f'</w:pPr></w:p>'
        ))
    for p in h2_elems:
        p.addnext(parse_xml(
            f'<w:p {nsdecls("w")}><w:pPr>'
            f'<w:pBdr><w:bottom w:val="single" w:sz="4" w:space="1" w:color="{NAVY}"/></w:pBdr>'
            f'<w:spacing w:before="0" w:after="60"/>'
            f'</w:pPr></w:p>'
        ))
    print(f"    {len(h1_elems)} gold + {len(h2_elems)} navy separators")


def _fix_korean_fonts(doc):
    from docx.oxml.ns import qn, nsdecls
    from docx.oxml import parse_xml

    CJK = [(0xAC00, 0xD7A3), (0x1100, 0x11FF), (0x3130, 0x318F),
           (0xA960, 0xA97F), (0xD7B0, 0xD7FF), (0x4E00, 0x9FFF)]

    def has_cjk(text):
        return text and any(any(s <= ord(c) <= e for s, e in CJK) for c in text)

    def fix(run):
        if not has_cjk(run.text):
            return
        rpr = run._r.get_or_add_rPr()
        rf = rpr.find(qn('w:rFonts'))
        if rf is None:
            rpr.insert(0, parse_xml(f'<w:rFonts {nsdecls("w")} w:eastAsia="{FONT_KR}"/>'))
        else:
            rf.set(qn('w:eastAsia'), FONT_KR)

    for p in doc.paragraphs:
        for r in p.runs:
            fix(r)
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for r in p.runs:
                        fix(r)


def _get_stats(doc):
    counts = {}
    for p in doc.paragraphs:
        sn = p.style.name if p.style else "None"
        counts[sn] = counts.get(sn, 0) + 1
    breaks = sum(1 for p in doc.paragraphs if p.paragraph_format.page_break_before)
    return {
        "paragraphs": len(doc.paragraphs),
        "tables": len(doc.tables),
        "h1": counts.get("Heading 1", 0),
        "h2": counts.get("Heading 2", 0),
        "h3": counts.get("Heading 3", 0),
        "h4": counts.get("Heading 4", 0),
        "code": sum(v for k, v in counts.items()
                    if any(kw in k for kw in ("Source Code", "Verbatim", "Code Block", "Preformatted"))),
        "blockquotes": counts.get("Body Text", 0),
        "breaks": breaks,
    }


def _print_stats(name, stats, output_file):
    est_pages = stats["breaks"] + stats["paragraphs"] // 45
    print(f"\n  {'='*50}")
    print(f"  {name}")
    print(f"  {'='*50}")
    print(f"  Paragraphs:  {stats['paragraphs']:,}")
    print(f"  Tables:      {stats['tables']:,}")
    print(f"  H1/H2/H3/H4: {stats['h1']}/{stats['h2']}/{stats['h3']}/{stats['h4']}")
    print(f"  Code blocks: {stats['code']}")
    print(f"  Blockquotes: {stats['blockquotes']}")
    print(f"  Page breaks: {stats['breaks']}")
    print(f"  Est. pages:  ~{est_pages}")
    print(f"  Output:      {output_file}")
    print(f"  {'='*50}")


# -- PDF Pipeline -------------------------------------------------------------
def find_xelatex():
    """Find XeLaTeX executable (TinyTeX or system install)."""
    paths = [
        os.path.join(os.environ.get("APPDATA", ""), "TinyTeX", "bin", "windows", "xelatex.exe"),
        r"C:\texlive\2024\bin\windows\xelatex.exe",
        r"C:\Program Files\MiKTeX\miktex\bin\x64\xelatex.exe",
    ]
    for p in paths:
        if os.path.exists(p):
            return p
    return "xelatex"


def run_pandoc_pdf(file_list, output_file):
    """Generate PDF via pandoc -> XeLaTeX pipeline."""
    pandoc = find_pandoc()
    xelatex = find_xelatex()

    filters = []
    if LUA_FILTER_DIAGRAM.exists():
        filters.extend(["--lua-filter", str(LUA_FILTER_DIAGRAM)])

    template_args = []
    if PDF_TEMPLATE.exists():
        template_args = [f"--template={PDF_TEMPLATE}"]

    cmd = [
        pandoc, "--from=markdown", "--to=pdf",
        f"--pdf-engine={xelatex}",
        *template_args,
        *filters,
        "--toc", "--toc-depth=3", "--standalone",
        "--wrap=preserve",
        f"--output={str(output_file)}",
    ]
    cmd.extend([str(f) for f in file_list])

    # Delete existing output so stale files don't mask failures
    if output_file.exists():
        try:
            output_file.unlink()
        except PermissionError:
            print(f"  [WARN] Cannot delete locked file: {output_file.name}")
            print(f"         Close the PDF viewer and retry.")
            return False

    # Run with explicit encoding and cwd to avoid temp path issues
    result = subprocess.run(cmd, capture_output=True, text=True,
                            encoding='utf-8', errors='replace',
                            cwd=str(PROJECT_ROOT))
    stderr = result.stderr or ""

    # Check if output was actually created
    if not output_file.exists() or output_file.stat().st_size == 0:
        error_lines = [l for l in stderr.split('\n')
                       if l.strip() and 'Hyper reference' not in l
                       and 'undefined references' not in l]
        print(f"  [FAIL] pandoc/xelatex error (exit {result.returncode}):")
        for line in error_lines[:30]:
            print(f"    {line}")
        return False

    # Success (warnings are expected from hyperref on single-pass)
    warn_count = stderr.count('WARNING')
    if warn_count:
        print(f"    ({warn_count} warnings, non-fatal)")
    return True


def generate_deliverable_pdf(key):
    """Generate a single deliverable as PDF."""
    cfg = DELIVERABLES[key]
    title = cfg["title"]
    output = PDF_OUTPUTS.get(key, OUTPUT_DIR / f"CheeseStock_{key}.pdf")

    print(f"\n{'='*60}")
    print(f"  {title} [PDF]")
    print(f"{'='*60}")

    file_list = build_file_list(cfg["files"])
    print(f"  Files: {len(file_list)}")
    for f in file_list:
        print(f"    {f.name}")

    if not file_list:
        print("  [SKIP] No files found")
        return

    print(f"  [1/1] pandoc -> XeLaTeX -> PDF...")
    if not run_pandoc_pdf(file_list, output):
        return

    size_kb = output.stat().st_size / 1024
    print(f"  [OK] {output.name} ({size_kb:.0f} KB)")


# -- Main ---------------------------------------------------------------------
def generate_deliverable(key):
    cfg = DELIVERABLES[key]
    output = cfg["output"]
    title = cfg["title"]

    print(f"\n{'='*60}")
    print(f"  {title}")
    print(f"{'='*60}")

    file_list = build_file_list(cfg["files"])
    print(f"  Files: {len(file_list)}")
    for f in file_list:
        print(f"    {f.name}")

    if not file_list:
        print("  [SKIP] No files found")
        return

    print(f"  [1/2] pandoc...")
    if not run_pandoc(file_list, output):
        return

    print(f"  [2/2] Post-processing...")
    stats = post_process(output)
    _print_stats(title, stats, output)


def main():
    # Parse args
    use_pdf = "--pdf" in sys.argv
    targets = []
    for arg in sys.argv[1:]:
        a = arg.lstrip("-")
        if a in DELIVERABLES:
            targets.append(a)
    if not targets:
        targets = ["exec", "tech", "theory", "d4"]

    mode = "PDF" if use_pdf else "DOCX"
    print("=" * 60)
    print(f"  CheeseStock ANATOMY V7 -- 4-Deliverable Generator [{mode}]")
    print("=" * 60)

    if use_pdf:
        # PDF pipeline: pandoc -> XeLaTeX -> PDF (no reference.docx needed)
        for key in targets:
            generate_deliverable_pdf(key)

        print(f"\n{'='*60}")
        print("  All done! Generated PDF files:")
        for key in targets:
            out = PDF_OUTPUTS.get(key, f"CheeseStock_{key}.pdf")
            print(f"    {out}")
    else:
        # DOCX pipeline: pandoc -> DOCX -> python-docx post-processing
        print("\n[Step 0] Reference template...")
        create_reference_template()

        for key in targets:
            generate_deliverable(key)

        print(f"\n{'='*60}")
        print("  All done! Generated DOCX files:")
        for key in targets:
            print(f"    {DELIVERABLES[key]['output']}")

    print(f"{'='*60}")


if __name__ == "__main__":
    main()
